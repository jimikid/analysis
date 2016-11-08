"""
Created on 02/24/2016, @author: sbaek
    - initial release
    V01 : 03/25/2016
     - def soft_index()
    V02 : 04/15/2016
     - values=OrderedDict().  dict() format change the orders and color codes on plots are not consistent. 
    V03 : 06/01/2016
     - plot() and plot_bar() are added
    V04 : 08/05/2016
     - move functions to eff_func.py
    V04 : 11/08/2016
     - if limit has one (xlimit, ylimit), then apply them into all plots, otherwise [ (x1 limit, y1 limit), (x2 limit, y2 limit) ..]

"""
from collections import OrderedDict
from os import makedirs, listdir
from os.path import dirname, exists

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib import cm


def plot(data=[], filename='fig',label=None, limit=None, fig_num=1, title='', marker='o-',
         grid=True, save=True, combine=False, hold=False,
         xtick=None, xtick_label=None, ytick=None, ytick_label=None,
         legend=None, figsize=(8, 8), scale=''):
    '''
    :param data: [ ([x1 values],[y1 values]), ([x2 values],[y2 values]) ..]
    :param label: [ (x1 label, y1 label), (x2 label, y2 label) ..] , if there is one set given, it is duplicated
    :param limit: [ (x1 limit, y1 limit), (x2 limit, y2 limit) ..]
    :param fig_num:
    :param title: ''
    :param marker:
    :param grid: Bool
    :param save: Bool
    :param combine: Bool  #plot datasets on one plot
    :param hold: Bool
    :param legend: []  # list of str names for multiple datasets on one plot
    :param figsize:
    :return:
    '''

    print '\n figure %s' % fig_num
    fig = plt.figure(fig_num, figsize=figsize)

    sp = len(data)

    cnt = 1
    for i in data:
        subplot = 100 * sp + 10 + cnt  #
        if combine: subplot = 111
        figure = fig.add_subplot(subplot)

        print ' plot %s' % filename
        try:
            #if len(legend) == len (data):
            figure.plot(i[0], i[1], marker, label=legend[cnt - 1])  # the number of legned has to to be same as plots
            figure.legend(loc='upper left')  # need to set the location.  label is given in plot()
        except:
            #else:
            figure.plot(i[0], i[1], marker)

        if grid: figure.grid(grid)

        if scale is 'log': figure.set_xscale("log")
        if scale is 'loglog':
            figure.set_xscale("log")
            figure.set_yscale("log")

        if label is not None:
            if len(label)==1:xlabel, ylabel = label[0][0], label[0][1]  #if one set is given duplicated
            else:xlabel, ylabel = label[cnt - 1][0], label[cnt - 1][1]
            plt.xlabel(xlabel)
            plt.ylabel(ylabel)

        if limit is not None:
            if isinstance(limit, tuple):
                #if limit has one (xlimit, ylimit), then apply them into all plots
                xlimit, ylimit = limit
            else:
                # otherwise,  limit has a list of tuples [(xlimit, ylimit), (xlimit, ylimit)..] , then apply them into each plots in order
                xlimit, ylimit = limit[cnt - 1][0], limit[cnt - 1][1]
                plt.xlim(xlimit)
                plt.ylim(ylimit)

        if grid: figure.grid(grid)
        plt.ticklabel_format(style='sci', axis='x', scilimits=(0,0))
        plt.ticklabel_format(style='sci', axis='y', scilimits=(0,0))

        if xtick is not None:plt.xticks(xtick, xtick_label)
        if ytick is not None: plt.yticks(ytick, ytick_label)

        print ' ...%s' % cnt
        if cnt == 1:
            plt.title(title)  # this has to be at the end. set title once on the top
        else: pass
        cnt = cnt + 1


    if hold: plt.show()  # show()hold plot.  cannot close.
    if save:
        fig_name = filename + '.png'
        fig.savefig(fig_name)
        print ' save %s \n' % (fig_name)
        plt.close(fig)
    else:pass

    return fig


def plot_bar(df, xticks, title='', save=True):
    '''
    - value referred to df.index.
    - file format
              flux_Sec_In_110_Bavg  flux_Sec_Out_110_Bavg
path_02_B []              0.442064               0.176952
path_04_B []              0.103560               0.256245

    :param df: list
    :param xticks: list
    :param title:
    :param save:
    :return:
    '''

    for i in df.keys():  # iteration at the number of keys
        labels=df.index.tolist()
        y_pos = np.arange(len(labels))
        plt.barh(y_pos, df[i], align='center', alpha=0.5)
        xticks = xticks
        plt.xticks(xticks)
        plt.yticks(y_pos, labels)
        plt.xlabel(title)
        # plt.title('')
    plt.show()

    if save:
        if not title:
            title = 'fig_bar'
        fig_name = title + '.png'
        plt.savefig(fig_name)
        print ' save %s' % (fig_name)
        plt.close()

def create_2plot_2yaxis(data1=[], data2=[], data3=[], data4=[], label1=['x','y1','y2'], label2=['x','y1','y2'], marker='-'):
    '''
    creat plots  with secondary y axis. Only work with 2 data   
    num : figure number,
    plot is adjusted with max and min values with scale
    '''
#    f, (ax1, ax2) = plt.subplots(1, 2, sharey=True)
    f, ax1 = plt.subplots(2)
    
    ax1[0].plot(data1[0][0], data1[0][1], 'b'+marker)
    ax1[0].plot(data1[0][0], data1[0][1], 'b'+marker)
    ax1[0].set_xlabel(label1[0])
    ax1[0].ticklabel_format(style='sci', axis='y', scilimits=(0,0))
    # Make the y-axis label and tick labels match the line color.
    ax1[0].set_ylabel(label1[1], color='b')
    
    ymin, ymax=plot_range(data1[0][1]) 
    ax1[0].set_ylim((ymin, ymax))
    
    ax2 = ax1[0].twinx()
    ax2.plot(data2[0][0], data2[0][1], 'g'+marker)  
    ax2.set_ylabel(label1[2], color='g')    
    ax2.ticklabel_format(style='sci', axis='y', scilimits=(0,0)) 
    
    ymin, ymax=plot_range(data2[0][1]) 
    ax2.set_ylim((ymin, ymax))   
    plt.ticklabel_format(style='sci', axis='x', scilimits=(0,0))
    plt.grid()  
  
    ax1[1].plot(data3[0][0], data3[0][1], 'b'+marker)
    ax1[1].plot(data3[0][0], data3[0][1], 'b'+marker)
    ax1[1].set_xlabel(label2[0])
    ax1[1].ticklabel_format(style='sci', axis='y', scilimits=(0,0))
    # Make the y-axis label and tick labels match the line color.
    ax1[1].set_ylabel(label2[1], color='b')
    
    ymin, ymax=plot_range(data3[0][1]) 
    ax1[1].set_ylim((ymin, ymax))
    
    ax2 = ax1[1].twinx()
    ax2.plot(data4[0][0], data4[0][1], 'g'+marker)  
    ax2.set_ylabel(label2[2], color='g')
    ax2.ticklabel_format(style='sci', axis='y', scilimits=(0,0))    

    ymin, ymax=plot_range(data4[0][1]) 
    ax2.set_ylim((ymin, ymax))
    plt.ticklabel_format(style='sci', axis='x', scilimits=(0,0))
    plt.grid()   


def create_2yaxis(data1=[], data2=[], fig_num=1, label=['x','y1','y2'], marker='-'):
    '''
    creat plots  with secondary y axis. Only work with 2 data
    num : figure number,
    plot is adjusted with max and min values with scale
    '''

    #fig, ax1 = plt.subplots(fig_num)
    fig, ax1 = plt.subplots(fig_num)

    ax1.plot(data1[0][0], data1[0][1], 'b'+marker)
    ax1.set_xlabel(label[0])
    ax1.ticklabel_format(style='sci', axis='y', scilimits=(0,0))
    # Make the y-axis label and tick labels match the line color.
    ax1.set_ylabel(label[1], color='b')


    ax2 = ax1.twinx()
    ax2.plot(data2[0][0], data2[0][1], 'g'+marker)
    ax2.set_ylabel(label[2], color='g')
    ax2.ticklabel_format(style='sci', axis='y', scilimits=(0,0))
    plt.ticklabel_format(style='sci', axis='x', scilimits=(0,0))

    plt.grid()

   
def create_path_dir(name='data', path = dirname(__file__)):
    '''
    creat a folder to save data under the present folder if it does not exist.
    and return the path in string form
    '''
    # the variables in definition are default, example          
    if not exists(path+'/'+name):
        print ' make dir'
        makedirs(path+'/'+name)
    newpath=path+'/'+name+'/'
    return newpath  
  


def search_csv(filename, dataset_path=''):    
    file_index = []
    #find out existing files and save index ti fuke_index,  check up to 500, 
    for i in range(10):
        if exists(dataset_path+filename+'000%s.csv' %i):
            file_index.append(i)    
    for i in range(10, 100):
        if exists(dataset_path+filename+'00%s.csv' %i):
            file_index.append(i)   

    for i in range(10):
        if exists(dataset_path+filename+'00%s.csv' %i):
            file_index.append(i)    
    for i in range(10, 100):
        if exists(dataset_path+filename+'0%s.csv' %i):
            file_index.append(i)   

    for i in range(10):
        if exists(dataset_path+filename+'0%s.csv' %i):
            file_index.append(i)                
    for i in range(10, 100):
        if exists(dataset_path+filename+'%s.csv' %i):
            file_index.append(i)   

    for i in range(10):
        if exists(dataset_path+filename+'%s.csv' %i):
            file_index.append(i)    
            
    return file_index 


def create_docx(path, key, filename='pics', width=None, fontsize=9):
    from docx import Document
    from docx.shared import Inches
    from docx.shared import Pt

    files=listdir(path)
    file_list=[]
    for i in files:
        if i.find(key) is not -1:
            print ' %s is found' %i
            file_list.append(i)

    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run()
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(fontsize)

    if width is None:
        width=5.5/len(file_list)

    for pic in file_list:
        short_path=path.split('/')[-3]+'/'+path.split('/')[-2]+'/'+path.split('/')[-1]
        run.add_text(short_path+pic+':  \n')

    para = doc.add_paragraph()
    run = para.add_run()
    for pic in file_list:
        run.add_text('  \n')
        run.add_picture(path+pic, width= Inches(width))
    doc.save_prjt('%s.docx' % filename)
    print ' save %s.docs' %filename
